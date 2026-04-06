# rag/kb_manager.py
"""
多知识库动态加载：按 kb_id 懒加载 document_store 与 retriever，支持多知识库并存。
非 ASCII 的 kb_id（如中文）会映射为 ASCII 存储目录，避免 FAISS 在 Windows 下无法写入路径。
"""
import json
import os
import threading
import uuid
from typing import Dict, Any, List

from config.settings import KB_ROOT, EMBEDDING_DIM
from config.runtime import CURRENT_KB

# 延迟导入 Haystack Document，避免在模块加载时触发 Pydantic v2 兼容性问题
# 将在需要时通过 _get_document_class() 获取
_Document = None

def _get_document_class():
    """延迟导入 Haystack Document 类"""
    global _Document
    if _Document is None:
        from haystack import Document
        _Document = Document
    return _Document


# 缓存：storage_key -> {"document_store": ..., "retriever": ...}
_kb_cache: Dict[str, Dict[str, Any]] = {}
_lock = threading.Lock()

_MAPPING_FILE = os.path.join(KB_ROOT, "_mapping.json")
_NUMERIC_PREFIX = "kb_"
_NUMERIC_WIDTH = 4  # kb_0001, kb_0002 ...


def _is_ascii_safe_id(kb_id: str) -> bool:
    """仅含 ASCII 字母数字、下划线、连字符时可直接作目录名。"""
    if not kb_id:
        return False
    return all(c.isascii() and (c.isalnum() or c in "_-") for c in kb_id)


def _default_mapping() -> Dict[str, Any]:
    return {"display_to_storage": {}, "storage_to_display": {}, "next_numeric_id": 1}


def _load_mapping() -> Dict[str, Any]:
    if not os.path.isfile(_MAPPING_FILE):
        return _default_mapping()
    try:
        with open(_MAPPING_FILE, "r", encoding="utf-8") as f:
            m = json.load(f)
        if not isinstance(m, dict):
            return _default_mapping()
        m.setdefault("display_to_storage", {})
        m.setdefault("storage_to_display", {})
        if "next_numeric_id" not in m or not isinstance(m.get("next_numeric_id"), int):
            m["next_numeric_id"] = 1
        m = _migrate_legacy_non_ascii_mappings(m)
        # 若迁移导致变更，写回文件（避免重复迁移）
        _save_mapping(m)
        return m
    except Exception:
        return _default_mapping()


def _save_mapping(m: Dict[str, Any]) -> None:
    os.makedirs(os.path.dirname(_MAPPING_FILE), exist_ok=True)
    with open(_MAPPING_FILE, "w", encoding="utf-8") as f:
        json.dump(m, f, ensure_ascii=False, indent=2)


def _is_numeric_storage_key(s: str) -> bool:
    if not s or not isinstance(s, str):
        return False
    if not s.startswith(_NUMERIC_PREFIX):
        return False
    tail = s[len(_NUMERIC_PREFIX) :]
    return bool(tail) and tail.isdigit()


def _allocate_numeric_storage_key(m: Dict[str, Any]) -> str:
    """
    分配一个新的数字 storage_key（kb_0001...），并自增 next_numeric_id。
    若对应目录已存在则继续自增直到找到空位（保证不会覆盖已有知识库）。
    """
    next_id = int(m.get("next_numeric_id") or 1)
    while True:
        key = f"{_NUMERIC_PREFIX}{next_id:0{_NUMERIC_WIDTH}d}"
        kb_dir = os.path.join(KB_ROOT, key)
        storage_to_display = m.get("storage_to_display") or {}
        if not os.path.exists(kb_dir) and key not in storage_to_display:
            m["next_numeric_id"] = next_id + 1
            return key
        next_id += 1


def _migrate_legacy_non_ascii_mappings(m: Dict[str, Any]) -> Dict[str, Any]:
    """
    将历史遗留的“中文展示名 -> kb_<md5>”映射迁移为“中文展示名 -> kb_0001”形式，并尽量重命名磁盘目录。
    仅迁移非 ASCII 安全的展示名，避免影响像 test1/zhongyi 这类原本就是目录名的知识库。
    """
    display_to_storage = m.get("display_to_storage") or {}
    storage_to_display = m.get("storage_to_display") or {}

    # 避免在遍历时修改 dict
    items = list(display_to_storage.items())
    changed = False

    for display_id, old_storage in items:
        if not display_id or _is_ascii_safe_id(str(display_id)):
            continue
        if not old_storage or _is_numeric_storage_key(str(old_storage)):
            continue

        # 为该 display_id 分配新的数字目录
        new_storage = _allocate_numeric_storage_key(m)

        old_dir = os.path.join(KB_ROOT, str(old_storage))
        new_dir = os.path.join(KB_ROOT, str(new_storage))

        # 尽量迁移现有数据：若旧目录存在且新目录不存在，做重命名
        try:
            if os.path.isdir(old_dir) and not os.path.exists(new_dir):
                os.rename(old_dir, new_dir)
        except Exception:
            # 迁移失败时不中断；仍更新映射，让后续新写入落到数字目录
            pass

        # 更新映射
        display_to_storage[str(display_id)] = str(new_storage)
        storage_to_display.pop(str(old_storage), None)
        storage_to_display[str(new_storage)] = str(display_id)
        changed = True

    if changed:
        m["display_to_storage"] = display_to_storage
        m["storage_to_display"] = storage_to_display
    return m


def get_storage_key(kb_id: str) -> str:
    """
    将用户传入的 kb_id 解析为磁盘上的 ASCII 目录名（storage_key）。
    - 若 kb_id 为空/default/haystack → 返回 CURRENT_KB。
    - 若 kb_id 为纯 ASCII 安全字符 → 直接转小写作为 storage_key。
    - 否则（含中文等）→ 用递增数字目录名 kb_0001... 并写入 _mapping.json（避免中文路径/编码问题）。
    """
    if not kb_id or not str(kb_id).strip():
        return CURRENT_KB
    kb_id = str(kb_id).strip()
    lower = kb_id.lower()
    if lower in ("default", "haystack"):
        return CURRENT_KB
    # 允许直接传 storage_key（例如前端缓存了 kb_0001）
    if _is_numeric_storage_key(kb_id):
        return kb_id
    if _is_ascii_safe_id(kb_id):
        return lower
    m = _load_mapping()
    display_to_storage = m.get("display_to_storage") or {}
    storage_to_display = m.get("storage_to_display") or {}
    if kb_id in display_to_storage:
        return display_to_storage[kb_id]
    key = _allocate_numeric_storage_key(m)
    display_to_storage[kb_id] = key
    storage_to_display[key] = kb_id
    m["display_to_storage"] = display_to_storage
    m["storage_to_display"] = storage_to_display
    _save_mapping(m)
    return key


def get_display_id(storage_key: str) -> str:
    """由 storage_key 得到对用户展示的 kb_id。"""
    if not storage_key:
        return storage_key
    m = _load_mapping()
    return m["storage_to_display"].get(storage_key, storage_key)


def list_storage_keys() -> List[str]:
    """列出 KB_ROOT 下所有用作知识库的目录名（即 storage_key），排除 _mapping 等。"""
    if not os.path.isdir(KB_ROOT):
        return []
    out = []
    for name in os.listdir(KB_ROOT):
        if name.startswith("_"):
            continue
        full = os.path.join(KB_ROOT, name)
        if os.path.isdir(full):
            out.append(name)
    return sorted(out)


def remove_mapping_for_display_id(kb_id: str) -> None:
    """删除知识库后从映射中移除该展示名。"""
    m = _load_mapping()
    if kb_id in m["display_to_storage"]:
        sk = m["display_to_storage"].pop(kb_id)
        m["storage_to_display"].pop(sk, None)
        _save_mapping(m)


def _resolve_kb_id(kb_id: str) -> str:
    """解析 kb_id 为 storage_key（用于所有磁盘路径），保持与 get_storage_key 一致。"""
    return get_storage_key(kb_id)


def _get_document_store_for_kb(kb_id: str):
    """为指定 kb_id 获取或创建 FAISS 文档存储。非 ASCII 名会映射为 ASCII 存储目录，避免 FAISS 无法写入。"""
    import faiss
    from haystack.document_stores import FAISSDocumentStore
    from storage.retriever import get_retriever as make_retriever

    storage_key = get_storage_key(kb_id)
    kb_dir = os.path.abspath(os.path.join(KB_ROOT, storage_key))
    db_file = os.path.join(kb_dir, "docs.db")
    base_faiss = os.path.join(kb_dir, "faiss.index")
    os.makedirs(kb_dir, exist_ok=True)

    possible_index_files = [
        base_faiss,
        base_faiss.replace(".index", ".faiss") if base_faiss.endswith(".index") else base_faiss + ".faiss",
        base_faiss if not base_faiss.endswith(".index") else base_faiss,
    ]
    if not base_faiss.endswith(".index") and not base_faiss.endswith(".faiss"):
        possible_index_files.append(base_faiss + ".index")

    found_index_file = None
    for p in possible_index_files:
        if os.path.exists(p):
            found_index_file = p
            break

    if found_index_file:
        try:
            index = faiss.read_index(found_index_file)
            if index.d == EMBEDDING_DIM:
                config_file = found_index_file.rsplit(".", 1)[0] + ".json"
                if os.path.exists(config_file):
                    try:
                        store = FAISSDocumentStore.load(found_index_file)
                        print(f"[KB={storage_key}] 已加载已有 FAISS 索引: {found_index_file}")
                        return store
                    except Exception as e:
                        # 典型场景：docs.db 中已有文档，但索引为空或不一致（例如报
                        # "The number of documents in the SQL database (...) doesn't match the number of embeddings in FAISS (...)")
                        # 此时忽略旧索引，落到下方分支按数据库内容重建索引。
                        print(f"[KB={storage_key}] 现有 FAISS 索引与数据库不一致，将丢弃并重建: {e}")
                # 无 config 时无法挂载已有 faiss，落到下方新建 store 并视情况重建索引
        except Exception as e:
            print(f"[KB={storage_key}] 加载索引文件失败，将创建新索引: {e}")

    store = FAISSDocumentStore(
        sql_url=f"sqlite:///{os.path.abspath(db_file).replace(chr(92), '/')}",
        embedding_dim=EMBEDDING_DIM,
        faiss_index_factory_str="Flat",
        # 初始化时先不强制检查“文档数 == 向量数”，避免
        # docs.db 已有文档但索引为空/不一致时直接抛错；我们会在下方按需重建索引。
        validate_index_sync=False,
    )
    db_count = len(store.get_all_documents())
    if db_count > 0:
        print(f"[KB={storage_key}] 数据库中有 {db_count} 个文档，缺少索引，正在生成...")
        retriever = make_retriever(store)
        store.update_embeddings(retriever)
        faiss_path = base_faiss if base_faiss.endswith(".index") else base_faiss + ".index"
        store.save(faiss_path)
    else:
        print(f"[KB={storage_key}] 初始化新的 FAISS 索引")
    return store


class KnowledgeBaseManager:
    """多知识库管理器：懒加载、缓存，不重复加载 FAISS。"""

    def get_retriever(self, kb_id: str):
        """
        根据 kb_id 获取检索器。若未加载则初始化并从 knowledge_bases/{kb_id}/ 读取 faiss.index、docs.db。
        未传或 default/haystack 时使用 CURRENT_KB。
        """
        resolved = _resolve_kb_id(kb_id)
        with _lock:
            if resolved not in _kb_cache:
                store = _get_document_store_for_kb(resolved)
                from storage.retriever import get_retriever
                retriever = get_retriever(store)
                _kb_cache[resolved] = {"document_store": store, "retriever": retriever}
            return _kb_cache[resolved]["retriever"]

    def list_documents(self, kb_id: str, limit: int = 20, offset: int = 0):
        """
        列出知识库中的文档（脱敏，不含 content 字段）。
        """
        resolved = _resolve_kb_id(kb_id)
        with _lock:
            if resolved not in _kb_cache:
                store = _get_document_store_for_kb(resolved)
                _kb_cache[resolved] = {"document_store": store}
            store = _kb_cache[resolved]["document_store"]
        
        # 获取所有文档并脱敏
        all_docs = store.get_all_documents()
        # 脱敏：移除 content 字段
        result = []
        for doc in all_docs[offset:offset + limit]:
            result.append({
                "id": doc.id,
                "meta": doc.meta,
                "content": None,  # 脱敏，不返回 content
            })
        return result

    def count_documents(self, kb_id: str) -> int:
        """
        统计知识库中的文档数量。
        """
        resolved = _resolve_kb_id(kb_id)
        with _lock:
            if resolved not in _kb_cache:
                store = _get_document_store_for_kb(resolved)
                _kb_cache[resolved] = {"document_store": store}
            store = _kb_cache[resolved]["document_store"]
        
        return len(store.get_all_documents())

    def add_file(self, kb_id: str, file):
        """
        上传并解析文件：支持 txt / json / jsonl / doc / docx。
        返回 task_id（用于查询导入状态）。
        """
        resolved = _resolve_kb_id(kb_id)
        print(f"[add_file] kb_id={kb_id}, resolved={resolved}")
        
        with _lock:
            if resolved not in _kb_cache:
                store = _get_document_store_for_kb(resolved)
                _kb_cache[resolved] = {"document_store": store}
            store = _kb_cache[resolved]["document_store"]
        
        # 读取文件内容
        file_content = file.file.read()
        file_type = file.content_type or ""
        file_name = file.filename or "unknown"
        print(f"[add_file] file_name={file_name}, file_type={file_type}, size={len(file_content)}")
        
        # 根据文件类型解析内容
        text = ""
        try:
            if file_type in ["text/plain", "application/plain"] or file_name.endswith(".txt"):
                text = file_content.decode("utf-8", errors="ignore")
            elif file_name.endswith(".jsonl"):
                # JSONL 文件：每行一个JSON对象
                print(f"[add_file] Processing JSONL file")
                lines = file_content.decode("utf-8", errors="ignore").strip().split("\n")
                texts = []
                for line in lines:
                    if line.strip():
                        try:
                            obj = json.loads(line)
                            # 提取文本内容（假设有 "text" 或 "content" 字段）
                            if isinstance(obj, dict):
                                content = obj.get("text") or obj.get("content") or str(obj)
                                texts.append(content)
                            else:
                                texts.append(str(obj))
                        except Exception as e:
                            print(f"[add_file] Failed to parse line: {e}")
                            texts.append(line)
                text = "\n".join(texts)
                print(f"[add_file] Parsed {len(texts)} lines from JSONL")
            elif file_type in ["application/json"] or file_name.endswith(".json"):
                try:
                    json_data = json.loads(file_content.decode("utf-8"))
                    if isinstance(json_data, list):
                        text = "\n".join([str(item) for item in json_data])
                    else:
                        text = str(json_data)
                except Exception:
                    text = file_content.decode("utf-8", errors="ignore")
            elif file_name.endswith(".docx"):
                import docx
                import io
                doc_obj = docx.Document(io.BytesIO(file_content))
                text = "\n".join([para.text for para in doc_obj.paragraphs])
            elif file_name.endswith(".doc"):
                # 旧版doc文件处理较复杂，暂时返回错误
                raise ValueError("旧版 .doc 文件暂不支持，请转换为 .docx")
            else:
                # 尝试直接解码
                text = file_content.decode("utf-8", errors="ignore")
        except Exception as e:
            print(f"[add_file] File parsing error: {e}")
            raise ValueError(f"文件解析失败: {e}")
        
        print(f"[add_file] Parsed text length: {len(text)}")
        
        # 生成文档
        doc_id = str(uuid.uuid4())
        meta = {"source": file_name, "file_type": file_type}
        Document = _get_document_class()
        doc = Document(content=text, id=doc_id, meta=meta)
        
        print(f"[add_file] Writing document to store...")
        # 写入文档
        store.write_documents([doc])
        
        print(f"[add_file] Updating embeddings...")
        store.update_embeddings(self.get_retriever(resolved))
        
        print(f"[add_file] Document added successfully, doc_id={doc_id}")
        return doc_id

    def add_text(self, kb_id: str, text: str):
        """
        上传文本内容：直接添加文本到知识库。
        返回 task_id（文档ID）。
        """
        resolved = _resolve_kb_id(kb_id)
        with _lock:
            if resolved not in _kb_cache:
                store = _get_document_store_for_kb(resolved)
                _kb_cache[resolved] = {"document_store": store}
            store = _kb_cache[resolved]["document_store"]
        
        # 生成文档
        doc_id = str(uuid.uuid4())
        meta = {"source": "manual_text", "type": "text"}
        Document = _get_document_class()
        doc = Document(content=text, id=doc_id, meta=meta)
        
        # 写入文档
        store.write_documents([doc])
        store.update_embeddings(self.get_retriever(resolved))
        
        return doc_id

    def add_document(self, kb_id: str, doc: Dict[str, Any]):
        """
        添加单条文档：doc 字段为完整文档对象（包含 content 和 meta）。
        返回文档ID。
        """
        resolved = _resolve_kb_id(kb_id)
        with _lock:
            if resolved not in _kb_cache:
                store = _get_document_store_for_kb(resolved)
                _kb_cache[resolved] = {"document_store": store}
            store = _kb_cache[resolved]["document_store"]
        
        # 生成文档ID（如果doc中没有提供）
        doc_id = doc.get("id") or str(uuid.uuid4())
        content = doc.get("content", "")
        meta = doc.get("meta", {})
        
        # 创建Document对象
        Document = _get_document_class()
        doc_obj = Document(content=content, id=doc_id, meta=meta)
        
        # 写入文档
        store.write_documents([doc_obj])
        store.update_embeddings(self.get_retriever(resolved))
        
        return doc_id

    def get_document(self, kb_id: str, doc_id: str):
        """
        获取知识库中的单个文档。
        """
        resolved = _resolve_kb_id(kb_id)
        with _lock:
            if resolved not in _kb_cache:
                store = _get_document_store_for_kb(resolved)
                _kb_cache[resolved] = {"document_store": store}
            store = _kb_cache[resolved]["document_store"]
        
        return store.get_document_by_id(doc_id)

    def update_document(self, kb_id: str, doc_id: str, doc):
        """
        更新知识库中的文档。
        """
        resolved = _resolve_kb_id(kb_id)
        with _lock:
            if resolved not in _kb_cache:
                store = _get_document_store_for_kb(resolved)
                _kb_cache[resolved] = {"document_store": store}
            store = _kb_cache[resolved]["document_store"]
        
        # 获取现有文档
        existing_doc = store.get_document_by_id(doc_id)
        if not existing_doc:
            raise ValueError(f"文档 {doc_id} 不存在")
        
        # 更新内容和元数据
        if isinstance(doc, dict):
            if "content" in doc:
                existing_doc.content = doc["content"]
            if "meta" in doc:
                existing_doc.meta.update(doc["meta"])
        else:
            if hasattr(doc, "content"):
                existing_doc.content = doc.content
            if hasattr(doc, "meta"):
                existing_doc.meta.update(doc.meta)
        
        # 更新文档
        store.update_documents([existing_doc])
        store.update_embeddings(self.get_retriever(resolved))
        
        return existing_doc.id

    def delete_document(self, kb_id: str, doc_id: str):
        """
        删除知识库中的文档。
        """
        resolved = _resolve_kb_id(kb_id)
        with _lock:
            if resolved not in _kb_cache:
                store = _get_document_store_for_kb(resolved)
                _kb_cache[resolved] = {"document_store": store}
            store = _kb_cache[resolved]["document_store"]
        
        store.delete_documents(ids=[doc_id])
        return True

    def search(self, kb_id: str, query: str, top_k: int = 10):
        """
        在知识库中搜索文档。
        """
        resolved = _resolve_kb_id(kb_id)
        retriever = self.get_retriever(resolved)
        
        docs = retriever.retrieve(query, top_k=top_k)
        
        # 脱敏处理
        result = []
        for doc in docs:
            result.append({
                "id": doc.id,
                "content": doc.content,
                "meta": doc.meta,
                "score": doc.score,
            })
        return result

    def create_knowledge_base(self, kb_name: str) -> str:
        """
        创建知识库目录：仅创建 knowledge_bases/{storage_key} 目录（非 ASCII 名会映射为 ASCII），不写入任何文档。
        若已存在则返回已存在的 storage_key。
        """
        from rag.kb_manager import get_storage_key
        
        # 生成 storage_key
        storage_key = get_storage_key(kb_name)
        
        # 创建物理目录
        kb_dir = os.path.join(KB_ROOT, storage_key)
        os.makedirs(kb_dir, exist_ok=True)
        
        return storage_key

    def delete_knowledge_base(self, kb_id: str) -> bool:
        """
        删除知识库目录：仅删除 knowledge_bases/{storage_key} 目录（如果有文档会一并删除）。
        """
        import shutil
        
        resolved = _resolve_kb_id(kb_id)
        kb_dir = os.path.join(KB_ROOT, resolved)
        
        if os.path.exists(kb_dir):
            shutil.rmtree(kb_dir)
            return True
        
        return False


# 单例，供 server 与后续模块使用
kb_manager = KnowledgeBaseManager()
