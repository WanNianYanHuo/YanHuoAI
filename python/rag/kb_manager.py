# rag/kb_manager.py
"""
多知识库动态加载: 按 kb_id 懒加载 document_store 和 retriever, 支持多知识库并存.
非 ASCII 的 kb_id (如中文) 会映射到 ASCII 存储目录, 避免 FAISS 在 Windows 下无法写入路径.
"""
import json
import os
import threading
import uuid
from typing import Dict, Any, List

from config.settings import KB_ROOT, EMBEDDING_DIM
from config.runtime import CURRENT_KB

# 延迟导入 Haystack Document, 避免在模块加载时触发 Pydantic v2 兼容性问题
# 将在需要时通过 _get_document_class() 获取
_Document = None

def _get_document_class():
    """延迟导入 Haystack Document 类"""
    global _Document
    if _Document is None:
        from haystack import Document
        _Document = Document
    return _Document


def _split_text(text: str, file_name: str = "", chunk_size: int = 0, chunk_overlap: int = 0) -> List[str]:
    """
    根据文件类型切割文本
    
    Args:
        text: 输入文本
        file_name: 文件名 (用于判断文件类型)
        chunk_size: 保留参数 (不再使用)
        chunk_overlap: 保留参数 (不再使用)
    
    Returns:
        切割后的文本块列表
    """
    if not text or not text.strip():
        return []
    
    chunks = []
    
    # 1. JSONL 文件：每个 JSON 对象为一个 chunk
    if file_name.endswith(".jsonl"):
        print(f"[split] Processing as JSONL")
        lines = text.strip().split("\n")
        for line in lines:
            line = line.strip()
            if not line:
                continue
            try:
                obj = json.loads(line)
                if isinstance(obj, dict):
                    content = obj.get("text") or obj.get("content") or str(obj)
                    if content and content.strip():
                        chunks.append(content.strip())
                else:
                    content = str(obj).strip()
                    if content:
                        chunks.append(content)
            except Exception as e:
                # 如果解析失败，整行作为一个 chunk
                if line.strip():
                    chunks.append(line.strip())
        print(f"[split] JSONL split into {len(chunks)} chunks")
        return chunks
    
    # 2. JSON 文件：每行为一个 chunk
    if file_name.endswith(".json") or text.strip().startswith("{"):
        print(f"[split] Processing as JSON")
        try:
            # 尝试解析为 JSON 对象
            obj = json.loads(text)
            if isinstance(obj, dict):
                # 如果是单个对象，提取 text 或 content 字段
                if "text" in obj:
                    content = obj["text"]
                elif "content" in obj:
                    content = obj["content"]
                else:
                    content = str(obj)
                
                # 按行分割
                for line in content.split("\n"):
                    line = line.strip()
                    if line:
                        chunks.append(line)
            elif isinstance(obj, list):
                # 数组：每个元素为一个 chunk
                for item in obj:
                    if isinstance(item, dict):
                        content = item.get("text") or item.get("content") or str(item)
                        if content and content.strip():
                            chunks.append(content.strip())
                    else:
                        content = str(item).strip()
                        if content:
                            chunks.append(content)
            else:
                # 其他类型，按行分割
                for line in str(obj).split("\n"):
                    line = line.strip()
                    if line:
                        chunks.append(line)
        except json.JSONDecodeError:
            # 解析失败，按行分割
            for line in text.split("\n"):
                line = line.strip()
                if line:
                    chunks.append(line)
        print(f"[split] JSON split into {len(chunks)} chunks")
        return chunks
    
    # 3. 文本文件：以换行为分隔，每行一个 chunk
    print(f"[split] Processing as plain text")
    for line in text.split("\n"):
        line = line.strip()
        if line:
            chunks.append(line)
    print(f"[split] Text split into {len(chunks)} chunks")
    
    return chunks


# 缓存: storage_key -> {"document_store": ..., "retriever": ...}
_kb_cache: Dict[str, Dict[str, Any]] = {}
_lock = threading.Lock()

_MAPPING_FILE = os.path.join(KB_ROOT, "_mapping.json")
_NUMERIC_PREFIX = "kb_"
_NUMERIC_WIDTH = 4  # kb_0001, kb_0002 ...


def _is_ascii_safe_id(kb_id: str) -> bool:
    """仅含 ASCII 字母数字, 下划线, 连字符时可直接作目录名"""
    if not kb_id:
        return False
    return all(c.isascii() and (c.isalnum() or c in "_-") for c in kb_id)


def _default_mapping() -> Dict[str, Any]:
    return {"display_to_storage": {}, "storage_to_display": {}, "next_numeric_id": 1}


def _load_mapping() -> Dict[str, Any]:
    if not os.path.isfile(_MAPPING_FILE):
        return _default_mapping()
    try:
        with open(_MAPPING_FILE, "r", encoding="utf-8") as f:
            m = json.load(f)
        if not isinstance(m, dict):
            return _default_mapping()
        m.setdefault("display_to_storage", {})
        m.setdefault("storage_to_display", {})
        if "next_numeric_id" not in m or not isinstance(m.get("next_numeric_id"), int):
            m["next_numeric_id"] = 1
        m = _migrate_legacy_non_ascii_mappings(m)
        _save_mapping(m)
        return m
    except Exception:
        return _default_mapping()


def _save_mapping(m: Dict[str, Any]) -> None:
    os.makedirs(os.path.dirname(_MAPPING_FILE), exist_ok=True)
    with open(_MAPPING_FILE, "w", encoding="utf-8") as f:
        json.dump(m, f, ensure_ascii=False, indent=2)


def _is_numeric_storage_key(s: str) -> bool:
    if not s or not isinstance(s, str):
        return False
    if not s.startswith(_NUMERIC_PREFIX):
        return False
    tail = s[len(_NUMERIC_PREFIX) :]
    return bool(tail) and tail.isdigit()


def _allocate_numeric_storage_key(m: Dict[str, Any]) -> str:
    """分配一个新的数字 storage_key (kb_0001...), 并自增 next_numeric_id."""
    next_id = int(m.get("next_numeric_id") or 1)
    while True:
        key = f"{_NUMERIC_PREFIX}{next_id:0{_NUMERIC_WIDTH}d}"
        kb_dir = os.path.join(KB_ROOT, key)
        storage_to_display = m.get("storage_to_display") or {}
        if not os.path.exists(kb_dir) and key not in storage_to_display:
            m["next_numeric_id"] = next_id + 1
            return key
        next_id += 1


def _migrate_legacy_non_ascii_mappings(m: Dict[str, Any]) -> Dict[str, Any]:
    """将历史遗留的映射迁移为数字形式."""
    display_to_storage = m.get("display_to_storage") or {}
    storage_to_display = m.get("storage_to_display") or {}
    items = list(display_to_storage.items())
    changed = False

    for display_id, old_storage in items:
        if not display_id or _is_ascii_safe_id(str(display_id)):
            continue
        if not old_storage or _is_numeric_storage_key(str(old_storage)):
            continue

        new_storage = _allocate_numeric_storage_key(m)
        old_dir = os.path.join(KB_ROOT, str(old_storage))
        new_dir = os.path.join(KB_ROOT, str(new_storage))

        try:
            if os.path.isdir(old_dir) and not os.path.exists(new_dir):
                os.rename(old_dir, new_dir)
        except Exception:
            pass

        display_to_storage[str(display_id)] = str(new_storage)
        storage_to_display.pop(str(old_storage), None)
        storage_to_display[str(new_storage)] = str(display_id)
        changed = True

    if changed:
        m["display_to_storage"] = display_to_storage
        m["storage_to_display"] = storage_to_display
    return m


def get_storage_key(kb_id: str) -> str:
    """将用户传入的 kb_id 解析为磁盘上的 ASCII 目录名 (storage_key)."""
    if not kb_id or not str(kb_id).strip():
        return CURRENT_KB
    kb_id = str(kb_id).strip()
    lower = kb_id.lower()
    if lower in ("default", "haystack"):
        return CURRENT_KB
    if _is_numeric_storage_key(kb_id):
        return kb_id
    if _is_ascii_safe_id(kb_id):
        return lower
    m = _load_mapping()
    display_to_storage = m.get("display_to_storage") or {}
    storage_to_display = m.get("storage_to_display") or {}
    if kb_id in display_to_storage:
        return display_to_storage[kb_id]
    key = _allocate_numeric_storage_key(m)
    display_to_storage[kb_id] = key
    storage_to_display[key] = kb_id
    m["display_to_storage"] = display_to_storage
    m["storage_to_display"] = storage_to_display
    _save_mapping(m)
    return key


def get_display_id(storage_key: str) -> str:
    """从 storage_key 得到对用户展示的 kb_id"""
    if not storage_key:
        return storage_key
    m = _load_mapping()
    return m["storage_to_display"].get(storage_key, storage_key)


def list_storage_keys() -> List[str]:
    """列出 KB_ROOT 下所有用作知识库的目录名 (即 storage_key), 排除 _mapping 等"""
    if not os.path.isdir(KB_ROOT):
        return []
    out = []
    for name in os.listdir(KB_ROOT):
        if name.startswith("_"):
            continue
        full = os.path.join(KB_ROOT, name)
        if os.path.isdir(full):
            out.append(name)
    return sorted(out)


def remove_mapping_for_display_id(kb_id: str) -> None:
    """删除知识库后从映射中移除该展示名"""
    m = _load_mapping()
    if kb_id in m["display_to_storage"]:
        sk = m["display_to_storage"].pop(kb_id)
        m["storage_to_display"].pop(sk, None)
        _save_mapping(m)


def _resolve_kb_id(kb_id: str) -> str:
    """解析 kb_id 到 storage_key (用于所有磁盘路径), 保持与 get_storage_key 一致"""
    return get_storage_key(kb_id)


def _get_document_store_for_kb(kb_id: str):
    """
    为指定 kb_id 获取或创建 FAISS 文档存储.
    Haystack 2.x 版本: 使用 haystack_integrations.document_stores.faiss
    """
    try:
        from haystack_integrations.document_stores.faiss import FAISSDocumentStore
    except ImportError:
        raise ImportError(
            "需要安装 faiss-haystack 包. 请运行: pip install faiss-haystack"
        )

    storage_key = get_storage_key(kb_id)
    
    # 使用缓存避免重复创建
    with _lock:
        if storage_key in _kb_cache:
            return _kb_cache[storage_key]["document_store"]
    
    kb_dir = os.path.abspath(os.path.join(KB_ROOT, storage_key))
    os.makedirs(kb_dir, exist_ok=True)

    index_path = os.path.join(kb_dir, "faiss_store")

    # 尝试加载已有索引
    if os.path.exists(f"{index_path}.faiss"):
        try:
            store = FAISSDocumentStore(
                index_path=index_path,
                embedding_dim=EMBEDDING_DIM
            )
            print(f"[KB={storage_key}] 已加载已有 FAISS 索引")
        except Exception as e:
            print(f"[KB={storage_key}] 加载索引失败: {e}")
            print(f"[KB={storage_key}] 尝试备份损坏的文件并创建新索引...")
            
            # 备份损坏的文件
            import time
            import shutil
            if os.path.exists(f"{index_path}.faiss"):
                backup_path = f"{index_path}_corrupted_{int(time.time())}.faiss"
                try:
                    shutil.move(f"{index_path}.faiss", backup_path)
                    print(f"[KB={storage_key}] 已备份损坏的索引文件到: {backup_path}")
                except Exception as backup_error:
                    print(f"[KB={storage_key}] 备份失败: {backup_error}")
            
            # 创建新的 document store
            store = FAISSDocumentStore(
                index_path=index_path,
                index_string="Flat",
                embedding_dim=EMBEDDING_DIM
            )
            print(f"[KB={storage_key}] 初始化新的 FAISS 索引")
    else:
        # 创建新的 document store
        store = FAISSDocumentStore(
            index_path=index_path,
            index_string="Flat",
            embedding_dim=EMBEDDING_DIM
        )
        print(f"[KB={storage_key}] 初始化新的 FAISS 索引")

    # 缓存 document store
    with _lock:
        _kb_cache[storage_key] = {"document_store": store}
    
    return store


def _safe_save_store(store, resolved_key: str, context: str):
    """
    安全保存当前知识库的 FAISS 索引到磁盘，带降级兜底，避免因 save 失败导致数据不落盘。
    context: 调用来源标记，便于日志排查（如 add_file / add_text 等）。
    """
    try:
        # 优先使用集成的 save 方法
        if hasattr(store, "save") and hasattr(store, "index_path"):
            store.save(store.index_path)
            print(f"[KB={resolved_key}] FAISS 索引已保存到磁盘 ({context}, 方法1)")
            return
        # 某些版本可能提供 save_to_disk
        if hasattr(store, "save_to_disk") and hasattr(store, "index_path"):
            store.save_to_disk(store.index_path)
            print(f"[KB={resolved_key}] FAISS 索引已保存到磁盘 ({context}, 方法2)")
            return
        print(f"[KB={resolved_key}] 警告：DocumentStore 不支持 save/index_path（{context}），可能无法持久化")
    except Exception as save_error:
        print(f"[KB={resolved_key}] 保存 FAISS 索引失败 ({context}): {save_error}")
        # 尝试使用底层 faiss 手动写索引
        try:
            import faiss  # type: ignore
            if hasattr(store, "index") and hasattr(store, "index_path"):
                faiss.write_index(store.index, f"{store.index_path}.faiss")
                # 手动写 json 元数据
                from pathlib import Path
                import json as _json
                path = Path(store.index_path)
                data = {
                    "documents": [doc.to_dict() for doc in store.documents.values()],
                    "id_map": store.id_map,
                    "inverse_id_map": store.inverse_id_map,
                    "next_id": store._next_id,
                }
                with open(path.with_suffix(".json"), "w", encoding="utf-8") as f:
                    _json.dump(data, f)
                print(f"[KB={resolved_key}] 使用 faiss.write_index 手动保存成功 ({context})")
        except Exception as manual_err:
            print(f"[KB={resolved_key}] 手动保存索引也失败 ({context}): {manual_err}")



class KnowledgeBaseManager:
    """多知识库管理器: 懒加载, 缓存, 不重复加载 FAISS. Haystack 2.x 版本"""

    def get_retriever(self, kb_id: str):
        """根据 kb_id 获取 document store (在 Haystack 2.x 中用于检索)."""
        resolved = _resolve_kb_id(kb_id)
        with _lock:
            if resolved not in _kb_cache:
                store = _get_document_store_for_kb(resolved)
                _kb_cache[resolved] = {"document_store": store}
            return _kb_cache[resolved]["document_store"]

    def list_documents(self, kb_id: str, limit: int = 20, offset: int = 0):
        """列出知识库中的文档 (脱敏, 不含 content 字段)"""
        resolved = _resolve_kb_id(kb_id)
        with _lock:
            if resolved not in _kb_cache:
                store = _get_document_store_for_kb(resolved)
                _kb_cache[resolved] = {"document_store": store}
            store = _kb_cache[resolved]["document_store"]

        # Haystack 2.x: filter_documents 不传 filters 参数时返回所有文档
        all_docs = store.filter_documents()

        result = []
        for doc in all_docs[offset:offset + limit]:
            result.append({
                "id": doc.id,
                "meta": doc.meta,
                "content": None,
            })
        return result

    def count_documents(self, kb_id: str) -> int:
        """统计知识库中的文档数量"""
        resolved = _resolve_kb_id(kb_id)
        with _lock:
            if resolved not in _kb_cache:
                store = _get_document_store_for_kb(resolved)
                _kb_cache[resolved] = {"document_store": store}
            store = _kb_cache[resolved]["document_store"]

        return store.count_documents()

    def add_file(self, kb_id: str, file):
        """上传并解析文件: 支持 txt / json / jsonl / doc / docx. 返回 task_id (文档ID)."""
        from storage.retriever import get_document_embedder

        resolved = _resolve_kb_id(kb_id)
        print(f"[add_file] kb_id={kb_id}, resolved={resolved}")

        with _lock:
            if resolved not in _kb_cache:
                store = _get_document_store_for_kb(resolved)
                _kb_cache[resolved] = {"document_store": store}
            store = _kb_cache[resolved]["document_store"]

        file_content = file.file.read()
        file_type = file.content_type or ""
        file_name = file.filename or "unknown"
        print(f"[add_file] file_name={file_name}, file_type={file_type}, size={len(file_content)}")

        text = ""
        try:
            if file_type in ["text/plain", "application/plain"] or file_name.endswith(".txt"):
                text = file_content.decode("utf-8", errors="ignore")
            elif file_name.endswith(".jsonl"):
                print(f"[add_file] Processing JSONL file")
                lines = file_content.decode("utf-8", errors="ignore").strip().split("\n")
                texts = []
                for line in lines:
                    if line.strip():
                        try:
                            obj = json.loads(line)
                            if isinstance(obj, dict):
                                content = obj.get("text") or obj.get("content") or str(obj)
                                texts.append(content)
                            else:
                                texts.append(str(obj))
                        except Exception as e:
                            print(f"[add_file] Failed to parse line: {e}")
                            texts.append(line)
                text = "\n".join(texts)
                print(f"[add_file] Parsed {len(texts)} lines from JSONL")
            elif file_type in ["application/json"] or file_name.endswith(".json"):
                try:
                    json_data = json.loads(file_content.decode("utf-8"))
                    if isinstance(json_data, list):
                        text = "\n".join([str(item) for item in json_data])
                    else:
                        text = str(json_data)
                except Exception:
                    text = file_content.decode("utf-8", errors="ignore")
            elif file_name.endswith(".docx"):
                import docx
                import io
                doc_obj = docx.Document(io.BytesIO(file_content))
                text = "\n".join([para.text for para in doc_obj.paragraphs])
            elif file_name.endswith(".doc"):
                raise ValueError("旧版 .doc 文件暂不支持, 请转换为 .docx")
            else:
                text = file_content.decode("utf-8", errors="ignore")
        except Exception as e:
            print(f"[add_file] File parsing error: {e}")
            raise ValueError(f"文件解析失败: {e}")

        print(f"[add_file] Parsed text length: {len(text)}")

        # 切割文档
        chunks = _split_text(text, file_name=file_name)
        print(f"[add_file] Split into {len(chunks)} chunks")

        if not chunks:
            raise ValueError("文档内容为空")

        # 创建多个文档
        Document = _get_document_class()
        docs = []
        for i, chunk in enumerate(chunks):
            doc_id = str(uuid.uuid4())
            meta = {
                "source": file_name, 
                "file_type": file_type,
                "chunk_index": i,
                "total_chunks": len(chunks)
            }
            doc = Document(content=chunk, id=doc_id, meta=meta)
            docs.append(doc)

        print(f"[add_file] Embedding {len(docs)} documents...")
        embedder = get_document_embedder()
        result = embedder.run(documents=docs)
        embedded_docs = result["documents"]

        print(f"[add_file] Writing {len(embedded_docs)} documents to store...")
        store.write_documents(embedded_docs)

        # 持久化当前知识库索引到磁盘，避免重启后数据丢失
        _safe_save_store(store, resolved, "add_file")
        
        print(f"[add_file] Documents added successfully, count={len(docs)}")
        return {"task_id": str(uuid.uuid4()), "doc_count": len(docs)}

    def add_text(self, kb_id: str, text: str):
        """上传文本内容: 直接添加文本到知识库. 返回文档ID"""
        from storage.retriever import get_document_embedder

        resolved = _resolve_kb_id(kb_id)
        with _lock:
            if resolved not in _kb_cache:
                store = _get_document_store_for_kb(resolved)
                _kb_cache[resolved] = {"document_store": store}
            store = _kb_cache[resolved]["document_store"]

        if not text or not text.strip():
            raise ValueError("文本内容不能为空")

        # 切割文本 (以换行为分隔)
        chunks = _split_text(text, file_name="")
        print(f"[add_text] Split into {len(chunks)} chunks")

        if not chunks:
            raise ValueError("文本内容为空")

        # 创建多个文档
        Document = _get_document_class()
        docs = []
        for i, chunk in enumerate(chunks):
            doc_id = str(uuid.uuid4())
            meta = {
                "source": "manual_text", 
                "type": "text",
                "chunk_index": i,
                "total_chunks": len(chunks)
            }
            doc = Document(content=chunk, id=doc_id, meta=meta)
            docs.append(doc)

        embedder = get_document_embedder()
        result = embedder.run(documents=docs)
        embedded_docs = result["documents"]

        store.write_documents(embedded_docs)

        # 持久化当前知识库索引到磁盘
        _safe_save_store(store, resolved, "add_text")
        return {"task_id": str(uuid.uuid4()), "doc_count": len(docs)}

    def add_document(self, kb_id: str, doc: Dict[str, Any]):
        """添加单条文档. 返回文档ID"""
        from storage.retriever import get_document_embedder
        from config.settings import CHUNK_SIZE, CHUNK_OVERLAP

        resolved = _resolve_kb_id(kb_id)
        with _lock:
            if resolved not in _kb_cache:
                store = _get_document_store_for_kb(resolved)
                _kb_cache[resolved] = {"document_store": store}
            store = _kb_cache[resolved]["document_store"]

        content = doc.get("content", "")
        if not content or not content.strip():
            raise ValueError("文档内容不能为空")

        # 切割文本
        chunks = _split_text(content, chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
        print(f"[add_document] Split into {len(chunks)} chunks")

        if not chunks:
            raise ValueError("文档内容为空")

        # 创建多个文档
        Document = _get_document_class()
        docs = []
        for i, chunk in enumerate(chunks):
            doc_id = doc.get("id") or str(uuid.uuid4())
            meta = doc.get("meta", {}).copy()
            meta.update({
                "chunk_index": i,
                "total_chunks": len(chunks)
            })
            doc_obj = Document(content=chunk, id=doc_id, meta=meta)
            docs.append(doc_obj)

        embedder = get_document_embedder()
        result = embedder.run(documents=docs)
        embedded_docs = result["documents"]

        store.write_documents(embedded_docs)

        # 持久化当前知识库索引到磁盘
        _safe_save_store(store, resolved, "add_document")
        return {"task_id": str(uuid.uuid4()), "doc_count": len(docs)}


    def get_document(self, kb_id: str, doc_id: str):
        """获取知识库中的单个文档"""
        resolved = _resolve_kb_id(kb_id)
        with _lock:
            if resolved not in _kb_cache:
                store = _get_document_store_for_kb(resolved)
                _kb_cache[resolved] = {"document_store": store}
            store = _kb_cache[resolved]["document_store"]

        # Haystack 2.x: 使用正确的过滤器格式
        docs = store.filter_documents(
            filters={"field": "id", "operator": "==", "value": doc_id}
        )
        return docs[0] if docs else None

    def update_document(self, kb_id: str, doc_id: str, doc):
        """更新知识库中的文档"""
        from storage.retriever import get_document_embedder

        resolved = _resolve_kb_id(kb_id)
        with _lock:
            if resolved not in _kb_cache:
                store = _get_document_store_for_kb(resolved)
                _kb_cache[resolved] = {"document_store": store}
            store = _kb_cache[resolved]["document_store"]

        # Haystack 2.x: 使用正确的过滤器格式
        existing_docs = store.filter_documents(
            filters={"field": "id", "operator": "==", "value": doc_id}
        )
        if not existing_docs:
            raise ValueError(f"文档 {doc_id} 不存在")

        existing_doc = existing_docs[0]

        if isinstance(doc, dict):
            if "content" in doc:
                existing_doc.content = doc["content"]
            if "meta" in doc:
                existing_doc.meta.update(doc["meta"])
        else:
            if hasattr(doc, "content"):
                existing_doc.content = doc.content
            if hasattr(doc, "meta"):
                existing_doc.meta.update(doc.meta)

        store.delete_documents([doc_id])

        embedder = get_document_embedder()
        result = embedder.run(documents=[existing_doc])
        embedded_docs = result["documents"]

        store.write_documents(embedded_docs)

        # 持久化当前知识库索引到磁盘
        try:
            if hasattr(store, "save") and hasattr(store, "index_path"):
                store.save(store.index_path)
                print(f"[KB={resolved}] 文档更新后 FAISS 索引已保存到磁盘 (update_document)")
        except Exception as save_error:
            print(f"[KB={resolved}] 保存 FAISS 索引失败 (update_document): {save_error}")
        return existing_doc.id

    def delete_document(self, kb_id: str, doc_id: str):
        """删除知识库中的文档"""
        resolved = _resolve_kb_id(kb_id)
        with _lock:
            if resolved not in _kb_cache:
                store = _get_document_store_for_kb(resolved)
                _kb_cache[resolved] = {"document_store": store}
            store = _kb_cache[resolved]["document_store"]

        store.delete_documents([doc_id])

        # 持久化当前知识库索引到磁盘
        _safe_save_store(store, resolved, "delete_document")
        return True

    def search(self, kb_id: str, query: str, top_k: int = 10):
        """在知识库中搜索文档"""
        from storage.retriever import get_text_embedder

        resolved = _resolve_kb_id(kb_id)
        store = self.get_retriever(resolved)

        embedder = get_text_embedder()
        result = embedder.run(text=query)
        query_embedding = result["embedding"]

        docs = store.search(
            query_embedding=query_embedding,
            top_k=top_k
        )

        result = []
        for doc in docs:
            result.append({
                "id": doc.id,
                "content": doc.content,
                "meta": doc.meta,
                "score": doc.score if hasattr(doc, "score") else None,
            })
        return result

    def create_knowledge_base(self, kb_name: str) -> str:
        """创建知识库目录"""
        storage_key = get_storage_key(kb_name)
        kb_dir = os.path.join(KB_ROOT, storage_key)
        os.makedirs(kb_dir, exist_ok=True)
        return storage_key

    def delete_knowledge_base(self, kb_id: str) -> bool:
        """删除知识库目录"""
        import shutil

        resolved = _resolve_kb_id(kb_id)
        kb_dir = os.path.join(KB_ROOT, resolved)

        with _lock:
            _kb_cache.pop(resolved, None)

        if os.path.exists(kb_dir):
            shutil.rmtree(kb_dir)
            return True

        return False


# 单例, 供 server 与后续模块使用
kb_manager = KnowledgeBaseManager()
